# -*- coding: utf-8 -*-
"""
YTS 合同生成模块
================
基于韩文合同模板（크리에이터 캠페인 진행 계약서.docx）一键生成合同。

流程（与产品约定）：
  1. 系统先把已知字段自动填好（网红名/金额/频道/平台/交付日/签署日=当日）
  2. 人在详情页核对、修改
  3. 确认后生成正式 Word 合同下载

约定：
  - 签署日期 = 生成当日
  - 交付日期（납품일/게시일）= 详情页「交稿截止」，可在生成表单里改
  - 网红个人信息（生日/地址/账户/税类型）一律留空，由网红本人填写
  - 条款正文 12 条原样保留，只替换动态字段
"""
import io
import os
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

TEMPLATE_NAME = "크리에이터 캠페인 진행 계약서.docx"
TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             TEMPLATE_NAME)


# ---------------------------------------------------------------------------
# 韩文大写金额（금액）转换
# ---------------------------------------------------------------------------
_DIGITS = ["", "일", "이", "삼", "사", "오", "육", "칠", "팔", "구"]
_UNITS = ["", "십", "백", "천"]
_BIG = ["", "만", "억", "조"]


def _four_digits(n: int) -> str:
    """把一个 0-9999 的组转成韩文大写。
    法律文书规范：leading 的 일 不省略（1500000 → 일백오십만，非 백오십만）"""
    if n == 0:
        return ""
    s = ""
    for i in range(3, -1, -1):
        d = (n // (10 ** i)) % 10
        if d == 0:
            continue
        s += _DIGITS[d] + _UNITS[i]
    return s


def hangul_amount(n) -> str:
    """整数金额 → 韩文大写。例：1500000 → 일백오십만, 1234567 → 일백이십삼만사천오백육십칠"""
    try:
        n = int(float(n))
    except (TypeError, ValueError):
        return ""
    if n == 0:
        return "영"
    groups = []
    while n > 0:
        groups.append(n % 10000)
        n //= 10000
    parts = []
    for i in range(len(groups) - 1, -1, -1):
        g = groups[i]
        if g == 0:
            continue
        s = _four_digits(g)
        if i > 0:
            s += _BIG[i]
        parts.append(s)
    return "".join(parts)


# ---------------------------------------------------------------------------
# docx 文本替换（处理跨 run 拆分，尽量保留格式）
# ---------------------------------------------------------------------------
def _replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    """在段落内替换 old→new，支持 old 被拆到多个 run 的情况。
    ⚠ 两个坑必须处理（否则如署名行 new 里包含 old 时会越替换越多）：
    1. 每次插入后把搜索游标移到新文本之后，避免 find 命中刚插入的内容；
    2. 兜底整段回写前，先挖掉已插入的 new 再判断是否真有残留。"""
    runs = paragraph.runs
    if not runs:
        return False
    combined = "".join(r.text for r in runs)
    if old not in combined:
        return False

    start = 0
    for _ in range(combined.count(old)):
        runs = paragraph.runs
        cur = "".join(r.text for r in runs)
        idx = cur.find(old, start)
        if idx < 0:
            break
        end = idx + len(old)
        pos, placed = 0, False
        for r in runs:
            r_start, r_end = pos, pos + len(r.text)
            pos = r_end
            if r_end <= idx or r_start >= end:
                continue
            pre = r.text[:max(0, idx - r_start)]
            post = r.text[max(0, min(len(r.text), end - r_start)):]
            if not placed:
                r.text = pre + new + post
                placed = True
            else:
                r.text = pre + post
        start = idx + len(new)  # 跳过刚插入的新文本

    # 兜底：极端 run 结构导致逐 run 没替干净。先把已插入的 new 挖掉再看
    # 是否仍有真残留，回写时只替换真残留（保留已插入的 new 原样）
    runs = paragraph.runs
    cur = "".join(r.text for r in runs)
    if old in cur.replace(new, "\x00"):
        fixed = new.join(seg.replace(old, new) for seg in cur.split(new))
        runs[0].text = fixed
        for r in runs[1:]:
            r.text = ""
    return True


def _iter_paragraphs(doc):
    """遍历正文 + 所有表格（含一层嵌套）里的段落。
    已实测：合并单元格不会造成段落重复访问（row.cells 重复返回的单元格
    内部段落对象各自独立），无需去重；切勿用 id() 去重——lxml 代理对象
    的 id 不稳定，会误跳段落"""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    for nt in cell.tables:
                        for nrow in nt.rows:
                            for ncell in nrow.cells:
                                for p in ncell.paragraphs:
                                    yield p


def _replace_doc(doc, old: str, new: str) -> int:
    """全文档替换，返回命中段落数"""
    n = 0
    for p in _iter_paragraphs(doc):
        if _replace_in_paragraph(p, old, new):
            n += 1
    return n


# ---------------------------------------------------------------------------
# 日期格式化（韩文，不补零，更符合法律文书习惯）
# ---------------------------------------------------------------------------
def _kor_date(d: datetime) -> str:
    return f"{d.year}년 {d.month}월 {d.day}일"


def _parse_date(s: str):
    """把 2026-08-21 / 2026.8.21 / 2026년8월21일 等解析成 datetime；失败返回 None"""
    if not s:
        return None
    s = str(s).strip().replace("년", "-").replace("월", "-").replace("일", "")
    s = s.replace(".", "-").replace("/", "-")
    parts = [p for p in s.split("-") if p.strip()]
    if len(parts) >= 3:
        try:
            return datetime(int(parts[0]), int(parts[1]), int(parts[2]))
        except ValueError:
            return None
    return None


# ---------------------------------------------------------------------------
# 主入口：生成合同
# ---------------------------------------------------------------------------
def generate_contract(fields: dict, template_path: str = None) -> bytes:
    """填充模板并返回合同 Word 的字节流。

    fields 需要：
      name          网红名（必填）
      amount        合同金额（韩币，必填）
      channel_url   频道链接（可空）
      delivery_date 交付/上线日期（可空，空则留模板占位）
      sign_date     签署日期（可空，默认当日）
    网红个人信息（生日/地址/账户/税类型）不填，留空给网红。
    """
    path = template_path or TEMPLATE_PATH
    if not os.path.exists(path):
        raise FileNotFoundError(f"合同模板不存在：{path}")
    doc = Document(path)

    name = str(fields.get("name") or "").strip()
    amount = fields.get("amount") or 0
    channel_url = str(fields.get("channel_url") or "").strip()

    sign_d = _parse_date(fields.get("sign_date")) or datetime.now()
    delivery_d = _parse_date(fields.get("delivery_date"))

    # ---- 1) 交付日期（납품일/게시일）优先替换 ----
    if delivery_d:
        _replace_doc(doc, "납품일 (게시일) : 2026년 MM월 DD일",
                     f"납품일 (게시일) : {_kor_date(delivery_d)}")

    # ---- 2) 签署日期（默认当日）：按段落特征定位——含 MM월 且不含
    #         납품일 的段落才是签署行（交付日在表格行里，带 납품일 前缀），
    #         三种空格形态都试（模板可能用普通/双/不换行空格）----
    sign_str = _kor_date(sign_d)
    for p in _iter_paragraphs(doc):
        t = "".join(r.text for r in p.runs)
        if "MM월" not in t or "납품일" in t:
            continue
        for pat in ("2026년\xa0 MM월\xa0 DD일",
                    "2026년  MM월  DD일",
                    "2026년 MM월 DD일"):
            if pat in t:
                _replace_in_paragraph(p, pat, sign_str)
                break

    # ---- 3) 合同金额（韩文大写 + 数字）----
    hangul = hangul_amount(amount)
    try:
        num = f"{int(float(amount)):,}"
    except (TypeError, ValueError):
        num = str(amount)
    _replace_doc(doc, "(₩ (금액))", f"(₩ {num})")
    _replace_doc(doc, "(금액)원정", f"{hangul}원정")

    # ---- 4) 网红名 / 频道 / 平台 ----
    # 首段当事方：크리에이터명(에이전시)
    _replace_doc(doc, "크리에이터명(에이전시)", name)
    # 캠페인 상세：참여 크리에이터
    _replace_doc(doc, "(크리에이터명 기재 , 채널URL)",
                 f"{name}, {channel_url}" if channel_url else name)
    # 合同名 & 캠페인 标题里的 (크리에이터명)
    _replace_doc(doc, "(크리에이터명)", name)
    # 平台：默认写 YouTube 영상（模板占位符句点后有两个空格，两种形态都试）
    _replace_doc(doc, "(플랫폼 기재.  예) 인스타그램 릴스)", "YouTube 영상")
    _replace_doc(doc, "(플랫폼 기재. 예) 인스타그램 릴스)", "YouTube 영상")
    # 签署栏：성명(에이전시명)
    _replace_doc(doc, "성명(에이전시명) :", f"성명(에이전시명) : {name}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def contract_filename(name: str) -> str:
    """生成合同文件名：계약서_网红名_日期.docx"""
    safe = "".join(ch for ch in str(name or "크리에이터")
                   if ch.isalnum() or ch in " _-").strip() or "크리에이터"
    return f"계약서_{safe}_{datetime.now().strftime('%Y%m%d')}.docx"
