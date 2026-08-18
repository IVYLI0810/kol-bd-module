#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
YTS Guide 生成模块

- 内置原版韩文 제작 가이드（来自 [YouTube Shopping]알리익스프레스 콘텐츠 제작 가이드.docx）
- 接千问 DashScope（OpenAI 兼容接口）生成「内容方向 & 强带货脚本建议」韩文章节
- 组装完整 guide（原版 + AI 章节），支持导出 Word（python-docx）

Key 读取顺序：环境变量/Secrets DASHSCOPE_API_KEY -> 本地 dashscope_key_local.py（不入库）
"""
import io
import os
import re

import requests

DASHSCOPE_URL = os.environ.get(
    "DASHSCOPE_URL",
    "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions")
MODEL = "qwen-plus"

# ---------------------------------------------------------------------------
# 原版 가이드（韩文，忠实于 Word 原件）
# ---------------------------------------------------------------------------
ORIGINAL_GUIDE_MD = """# [YouTube Shopping] 알리익스프레스 콘텐츠 제작 가이드

## 1. 콘텐츠
- **내용**: 방향성에 맞게 제품을 직접 선정한 후, 각 제품의 특징과 추천 이유가 자연스럽게 포함되도록 구성
- **방향성**: 애용템/추천템 소개 · 언박싱 및 제품 리뷰 · 일상 속 제품 활용 꿀팁
- **제품**: 2-5개 희망. 유튜브 쇼핑 태그 등록 상품 중 희망 상품 직접 선정 → 희망 상품 리스트 전달 후 알리 검수 진행 → 최종 확정된 상품 직접 구매 후 콘텐츠 제작
- **형식**: 쇼츠 or 롱폼（쇼츠 15초 이상, 롱폼 3분 이상）

## 2. 코드 혜택 (Code: YTSKOL)
- **이용 안내**: $10 이상 구매 시 $2 할인 받기
- **유효 기간**: KST 7/21 00:00 - 9/30 23:59

## 3. 심의 안내
- 콘텐츠 제작 완료 후, 업로드 전 **반드시 영상을 전달** 부탁드립니다.
- 내부 심의 확인 항목: 광고 관련 법규 준수 여부, 상품 정보 및 혜택 정보의 정확성 등
- 수정이 필요한 경우 피드백 전달 → 수정 완료 후 재심의 진행
- **힌트**: 보다 원활한 심의를 위해 영상 내에 타 플랫폼의 명칭이 포함된 화면이 노출되지 않도록 유의해 주세요:)

## 4. 필수 사항
- 유튜브 **쇼핑 태그 기능 활용** 필수
- 유튜브 **유료광고 표기** 필수
- **코드 혜택 정보 포함** 필수（코드: YTSKOL / 유효 기간 / 이용 안내）

## 5. 참고 사항
- **제목**: 제품 선정 후 자유롭게 구성 가능
- **썸네일**: 제품과 사용 장면이 잘 드러나는 방향 권장
- **정산 방식**: 판매 수익은 유튜브를 통해 직접 정산 / 제작비는 별도 지급
- **참고 영상**: 하봄 https://youtube.com/shorts/xIetj_6u_uo · 푸짐스 https://youtube.com/shorts/S3BT8PiziC0 · 켈리아 https://www.youtube.com/shorts/W6bs2y-ab70

## 6. 예상 진행 일정（전체 약 1개월）
- ① 협업 의향 확정（협업 여부 및 예상 업로드 일정 확인）
- ② 협업 확정 후 5일 이내: 희망 상품 리스트 및 선정 이유 전달 → 내부 확인 후 상품 구매 진행
- ③ 제품 수령 완료 후 수령 여부 공유
- ④ 제품 수령 후 3~5일 이내: 기획안 전달 및 피드백
- ⑤ 기획안 확정 후 촬영 및 편집 진행
- ⑥ 매월 16일 이전: 업로드 예정 영상 전달
- ⑦ 내부 심의（영업일 기준 약 3일）→ 수정 필요 시 피드백 전달
- ⑧ 심의 통과 시 21일 업로드 / 미통과 시 수정본 재심의（약 1영업일）
- ⑨ 매월 21일 영상 업로드
- ※ 제작비는 영상 업로드 후 대행사를 통해 별도 계약 및 정산 진행 예정
"""

AI_SECTION_TITLE = "## 7. 콘텐츠 방향 & 강전환 스크립트 제안 (AI)"

SYSTEM_PROMPT = (
    "당신은 AliExpress 한국 YouTube Shopping 프로젝트의 시니어 콘텐츠 디렉터입니다. "
    "기존 제작 가이드 뒤에 덧붙일 「콘텐츠 방향 & 강전환(강한 판매 유도) 스크립트 제안」 섹션을 한국어로 작성합니다. "
    "규칙: 1) 방향은 반드시 강전환 지향 — 제품 셀링포인트, 사용 장면, 구매 동기 부여에 집중. "
    "2) 스크립트는 쇼츠 타임라인 형식(0-3초 훅, 3-10초 페인포인트/장면, 제품 등장+셀링포인트, 사용 효과, 코드 혜택, CTA). "
    "3) 할인 코드 YTSKOL($10 이상 구매 시 $2 할인)과 유튜브 쇼핑 태그 클릭 유도를 자연스럽게 포함. "
    "4) 심의 규칙 준수: 유료광고 표기 안내, 타 플랫폼 명칭 노출 금지. "
    "5) 한국어로만 출력, 마크다운 형식(## / - / **굵은글자**) 사용, 서론 없이 바로 본문 출력."
)

USER_PROMPT_TPL = """인플루언서 정보:
-昵称(닉네임): {name}
- 垂类(카테고리): {category}
- 粉丝数(구독자): {subscribers}
- 选品清单(선정 희망 상품 링크):
{products}
- 计划上线(업로드 예정): {plan_month}

추가 요청 사항:
{requirements}

아래 구조대로 출력하세요:
## A. 콘텐츠 방향 제안（강전환 지향, 2-3개）
## B. 쇼츠 스크립트 제안（타임라인 형식）
## C. 전환 유도 화술 & 코드 혜택(YTSKOL) 안내 멘트
## D. 제목 / 해시태그 제안"""


def get_api_key() -> str:
    key = os.environ.get("DASHSCOPE_API_KEY", "").strip()
    if key:
        return key
    try:
        from dashscope_key_local import DASHSCOPE_API_KEY  # noqa: 本地调试用，不入库
        return (DASHSCOPE_API_KEY or "").strip()
    except ImportError:
        return ""


def build_prompt(collab: dict, requirements: str = "") -> list:
    products = collab.get("product_list") or []
    user = USER_PROMPT_TPL.format(
        name=collab.get("name") or "-",
        category=collab.get("category") or "-",
        subscribers=collab.get("subscribers") or "-",
        products="\n".join(f"  - {p}" for p in products) or "  - (아직 없음, 일반적인 강전환 방향으로 제안)",
        plan_month=collab.get("plan_month") or "-",
        requirements=requirements.strip() or "(없음, 기본 강전환 방향)",
    )
    return [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": user},
    ]


def call_dashscope(messages: list, timeout: int = 120) -> str:
    """调用千问，返回生成文本；无 key / 出错时抛 RuntimeError(友好文案)"""
    key = get_api_key()
    if not key:
        raise RuntimeError(
            "未配置 DASHSCOPE_API_KEY · DASHSCOPE_API_KEY가 설정되지 않았습니다. "
            "请在主站 Cloud Secrets 添加 DASHSCOPE_API_KEY（百炼控制台获取）")
    resp = requests.post(
        DASHSCOPE_URL,
        headers={"Authorization": f"Bearer {key}",
                 "Content-Type": "application/json"},
        json={"model": MODEL, "messages": messages, "temperature": 0.8},
        timeout=timeout,
    )
    if resp.status_code != 200:
        raise RuntimeError(f"千问接口返回 {resp.status_code}: {resp.text[:300]}")
    data = resp.json()
    try:
        return data["choices"][0]["message"]["content"].strip()
    except (KeyError, IndexError) as e:
        raise RuntimeError(f"千问返回结构异常: {str(data)[:300]}") from e


def assemble_full_guide(script_md: str) -> str:
    return ORIGINAL_GUIDE_MD.rstrip() + "\n\n" + AI_SECTION_TITLE + "\n\n" + script_md.strip() + "\n"


# ---------------------------------------------------------------------------
# Markdown(子集) -> Word
# ---------------------------------------------------------------------------
_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _add_para_with_bold(doc, text, style=None):
    para = doc.add_paragraph(style=style)
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        para.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])
    return para


def md_to_docx(md: str, title: str = "") -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal.font.size = Pt(10.5)
    for line in md.splitlines():
        s = line.rstrip()
        if not s.strip():
            continue
        if s.startswith("### "):
            _add_para_with_bold(doc, s[4:], style="Heading 3")
        elif s.startswith("## "):
            _add_para_with_bold(doc, s[3:], style="Heading 2")
        elif s.startswith("# "):
            _add_para_with_bold(doc, s[2:], style="Heading 1")
        elif s.lstrip().startswith("- "):
            _add_para_with_bold(doc, s.lstrip()[2:], style="List Bullet")
        else:
            _add_para_with_bold(doc, s)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
